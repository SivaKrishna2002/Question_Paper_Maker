from tkinter import *
from tkinter import ttk
from tkinter import messagebox
import docx
from tkinter import filedialog
from docx.enum.style import WD_STYLE_TYPE
from docx.shared import Pt, RGBColor, Inches ,Cm
from docx.oxml.ns import qn

from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.text import WD_LINE_SPACING


doc = docx.Document()

sections = doc.sections
for section in sections:
    section.top_margin = Cm(1)
    section.bottom_margin = Cm(0.5)
    section.left_margin = Cm(1.5)
    section.right_margin = Cm(1.5)



plan=Tk()
plan.geometry("900x500")
plan.title("Question paper MAKER")
plan.configure(bg="whitesmoke")
plan.resizable(0,0)






#=========================================window 1===================================================================
#=============================Verify and Submit Button=============================================================
def button1():
    a=cn.get()
    mnr=a.upper()
    b=sc.get()
    c=sn.get()
    d=mm.get()
    e=dr.get()
    f=cl.get()
    g=dy.get()
    h=n.get()
    # if(a=='' | b=='' | c=='' | d=='' | e==''| f==''| g==''):
    #     messagebox.showerror("Error","All fields are required!!")

    #print(a)
    # print(b)
    # print(c)
    # print(d)
    # print(e)
    # print(f)
    # print(g)
    # print(h)

    header_text1=doc.add_heading("Hall Ticket Number")
    header_text1.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    header_text1.paragraph_format.space_before = Pt(3)
    header_text1.paragraph_format.space_after = Pt(5)
    header_text1.add_run(" :                                                                                               C")
    # header_text1.add_run(".                                           .")
    header_text1.add_run("ode No: ")
    header_text1.add_run(b)

    head=doc.add_heading(mnr,1)

    head.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    head.paragraph_format.space_before = Pt(3)
    head.paragraph_format.space_after = Pt(5)
    head.line_spacing = 0.5

    header_text=doc.add_heading(h,1)
    header_text.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    header_text.paragraph_format.space_before = Pt(3)
    header_text.paragraph_format.space_after = Pt(5)
    header_text.line_spacing = 0.5
    header_text.add_run("-")
    header_text.add_run(g)

    header_text=doc.add_heading(c,1)
    header_text.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    header_text.paragraph_format.space_before = Pt(3)
    header_text.paragraph_format.space_after = Pt(5)
    header_text.line_spacing = 0.5
    header_text.add_run("-")
    header_text.add_run(f)

    header_text3=doc.add_heading("Time: ",1)
    header_text3.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    header_text3.paragraph_format.space_before = Pt(3)
    header_text3.paragraph_format.space_after = Pt(5)
    header_text3.add_run(e)
    header_text3.add_run("h")
    header_text3.add_run("r                                                                                                               M")
    header_text3.add_run("ax Mark : ")
    header_text3.add_run(d)
    header_text3.add_run(" M")

    newline=doc.add_paragraph("______________________________________________________________________________________________________________________________")

    b1.config(state=NORMAL)
    b2.config(state=NORMAL)
    b3.config(state=NORMAL)
def saveas():
    file=filedialog.asksaveasfilename(title="Save As",initialfile=".docx",defaultextension=".0docx",filetypes=[("word files","*.docx"),("All files","*.*")])
    doc.save(file)

#===================================REfresh button========================================================================
def refresh():
    b1["state"]=NORMAL
    b2["state"]=NORMAL
#========================================================================================================================================================

#===================================Discriptive button===========================================================================
def discriptive():

    #===============================Click button in discriptive======================================================
    def click():
        #====================================== Window-2 ===================================================================
        win=Toplevel(plan)
        win.resizable(0,600)
        win.title("Question paper maker")
        l=['a','b','c','d','e','f','g','h','i','j','k']
        win.geometry("300x300")
        n=noofsec.get()
        lis=[]
        r=1
        #=========================== Section submit ==========================================================================
        def sectionsubmit():
            for i in range(1,n+1):
                if(l[i].get()==''):
                        messagebox.showerror("Error","All fields are required!!")
                p=int(l[i].get())
                lis.append(p)
            print(lis)
            
            #================================ Window-3 ===================================================
            def data():
                
                heading=[]
                questions=[]
                c=1
                l1=Label(frame,text="SECTIONS BLOCK INFORMATION",bg="skyblue",fg="red",font="arial 13 bold",justify=CENTER).grid(row=0,column=1,pady=10,padx=130,columnspan=100)
                #=============================== Window-3 submit =============================================================
                def sub():
                    saveas.config(state=NORMAL)
                    x=0 
                    po=0
                    for i in range(0,n):
                        win.destroy()
                        print(heading[i].get())
                        doc.add_heading(heading[i].get(),2)
                        # doc.space_after = Pt(2)
                        paragraph1=doc.add_paragraph(mam[po].get())


                        paragraph1.add_run("X")
                        po=po+1

                        paragraph1.add_run(mam[po].get())
                        paragraph1.add_run("=")
                        po=po+1

                        paragraph1.add_run(mam[po].get())
                        paragraph1.add_run("M")
                        po=po+1
                        paragraph1.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

                        for j in range(0,lis[i]):
                            print(questions[x].get())
                            ques=doc.add_paragraph(questions[x].get())
                            ques.style="List Number"
                            x=x+1
                    frame.destroy()

                #s=['aa','bb','cc','dd','ee','ff','gg','hh','ii','jj','kk','ll','mm','nn','oo','pp','qq','rr','ss','tt','uu','vv','ww','xx','yy','zz']  
                #z=0
                q=65
                mam=[]
                for i in range(0,n):
                
                    f=chr(q) 
                    l1=Label(frame,font="arial 10 bold",relief="groove",text="Heading of SECTION {} :".format(i+1)).grid(row=c,column=1,pady=5,padx=5)
                    lp=Label(frame,text="                                      ",font="arial 15 bold").grid(row=c,column=2)
                    f=StringVar()
                    entry123=Entry(frame,width=5,justify=CENTER,border=3,textvariable=f).grid(row=c,column=3)
                    mam.append(f)
                    q=q+1
                    h=chr(q)
                    lp1=Label(frame,text="x",font="arial 15 bold",border=3).grid(row=c,column=4)

                    h=StringVar()
                    entry456=Entry(frame,width=5,justify=CENTER,border=3,textvariable=h).grid(row=c,column=5)
                    mam.append(h)
                    q=q+1
                    w=chr(q)

                    lp=Label(frame,text="=",font="arial 15 bold").grid(row=c,column=6)
                    w=StringVar()
                    entry456=Entry(frame,width=5,justify=CENTER,border=3,textvariable=w).grid(row=c,column=7)
                    mam.append(w)
                    q=q+1
                    l[i]=StringVar()
                    p=l[i]
                    c=c+1
                    heading.append(p)
                    print(heading)
                    e=Entry(frame,width=90,textvariable=l[i],relief="groove",border=5).grid(row=c,column=1,columnspan=100,padx=10,pady=10)

                    for j in range(0,lis[i]):
                        c=c+1
                        l1=Label(frame,fg="blue",text="Question {}:".format(j+1)).grid(row=c,column=1)
                        c=c+1
                        l[j]=StringVar()
                        e=Entry(frame,width=90,textvariable=l[j],border=1).grid(row=c,column=1,columnspan=100,pady=5)
                        c=c+1
                        z=l[j]
                        questions.append(z)
                        print(questions)
                print(mam)    
                bhgv1=Button(frame,width=15,text="Submit",command=sub,fg="white",bg="red").grid(row=c,column=1,pady=30,padx=20)
   
    
        
            def myfunction(event):
                canvas.configure(scrollregion=canvas.bbox("all"),width=570,height=570)

            root=Toplevel(win)
            # sizex = 570
            # sizey = 650
            # posx  = 100
            # posy  = 100
            # root.wm_geometry("%dx%d+%d+%d" % (sizex, sizey, posx, posy))
            root.geometry("610x600")
            root.resizable(0,0)

            myframe=Frame(root,relief=SOLID,width=50,height=100,bd=1)
            myframe.place(x=10,y=10)
            canvas=Canvas(myframe)
            frame=Frame(canvas)
            myscrollbar=Scrollbar(myframe,orient="vertical",command=canvas.yview)
            canvas.configure(yscrollcommand=myscrollbar.set)
            myscrollbar.pack(side="right",fill="y")
            canvas.pack(side="left")
            canvas.create_window((0,0),window=frame,anchor='nw')
            frame.bind("<Configure>",myfunction)
            data()
            root.mainloop()


    
            #===========================================================================================================
        l1=Label(win,text="SECTIONS BLOCK",bg="skyblue",fg="red",font="arial 13 bold",justify=CENTER).grid(row=0,column=1,pady=10,padx=70,columnspan=100)
        for i in range(1,n+1):
            l3=Label(win,font="arial 10 bold",text="Count of Questions in SECTION {}:".format(i)).grid(row=r,column=1,pady=10,columnspan=20)
            r=r+1
            l[i]=StringVar()
            e=Entry(win,textvariable=l[i],width=40,justify=CENTER,font="arial 10 bold",fg="red").grid(row=r,column=1,padx=3,columnspan=50)
            r=r+1
        b1=Button(win,text="Submit",command=sectionsubmit,bg="red",fg="white").grid(row=r,column=2,pady=10,padx=20)
        r=r+1
        l8=Label(win,text="*fill all the fields",fg="red").grid(row=r,column=1)


        win.mainloop()
        #==========================================================================================================

    l=Label(plan,text="Total No of Sections").place(x=30,y=390)
    noofsec=IntVar()
    ee=Entry(plan,width=50,textvariable=noofsec).place(x=30,y=420)
    btn1=Button(plan,text="Next",fg="red",bg="white",border=2,font='camble 10 bold',command=click).place(x=120,y=450)
    b2.config(state = DISABLED)
#==================================== MCQS =============================================================
def mcqs():
    
    def check():
        def myfunction(event):
            canvas.configure(scrollregion=canvas.bbox("all"),width=760,height=580)
    
        def next(box):
            box.tkraise()

        chos=Toplevel()
        chos.rowconfigure(0,weight=1)
        chos.columnconfigure(0,weight=1)
        
        sizex = 800
        sizey = 600
        posx  = 100
        posy  = 100
        chos.wm_geometry("%dx%d+%d+%d" % (sizex, sizey, posx, posy))
        chos.resizable(0,0)

        def done():
            y=nwe.get()
            lis=[]
            q=nwr.get()
            # doc=docx.Document()
            # sections = doc.sections
            # for section in sections:
            #     section.top_margin = Cm(1)
            #     section.bottom_margin = Cm(0.5)
            #     section.left_margin = Cm(1.5)
            #     section.right_margin = Cm(1.5)

            nr=doc.add_heading("I . CHOOSE THE CORRECT OPTIONS FROM THE FOLLOWING QUESTIONS :",1)


            ans=y*q
            m=str(ans)
            p="MAX MARKS : "+m+" M"



            def donep():
                pass
            
            
            

            def saveme():
                nrw=doc.add_heading(p,2)
                nrw.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
                doc.add_paragraph("                   ")
                




                for i in lis:
                    for j in i:
                        if(j==i[0]):
                            qs=doc.add_paragraph(j.get())
                            qs.style="List Number"
                            nrtw=doc.add_heading("  [   ]  ",2)
                            nrtw.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
                            opr=doc.add_paragraph("           A) ")
                        if(j==i[1]):
                            opr.add_run(j.get())
                            opr.add_run("                                           B) ")
                        if(j==i[2]):
                            opr.add_run(j.get())
                            oprd=doc.add_paragraph("          C) ")
                        if(j==i[3]):
                            oprd.add_run(j.get())

                            oprd.add_run("                                          D) ")
                        if(j==i[4]):
                            oprd.add_run(j.get())
                        print("---------------------------------")
                        print(j.get())







                saveas.config(state=NORMAL)
                chos.destroy()

                # doc.save("newtuyr.docx")

            wins=Frame(frame)


            btr=Button(wins,text="<-- BACK",bg="red",fg="white",font=("Timens new roman",10,"bold"),command=lambda : next(winr)).grid(row=0,column=0)

            lard=Label(wins,text="MCQ'S QUESTIONS",font=("Timens new roman",18,"underline"),fg="red").grid(row=0,column=2,columnspan=100)

            wins.grid(row=0,column=0,sticky="nsew")
            c=1


            for i in range(1,y+1):
                qt=Label(wins,font = ("Times new roman",13,"bold"),text = ("Question {} : ".format(i+1))).grid(row=c,column=0,pady=15)
                k="text"+str(i)
                k=StringVar()
                en1=Entry(wins,width=100,relief="solid",bd=1.5,textvariable=k).grid(row=c,column=1,columnspan=100)
                c=c+1
                qtr=Label(wins,font = ("Times new roman",13,"bold"),text =("A).")).grid(row=c,column=1,pady=10)
                a="opna"+str(i)
                a=StringVar()

                eny=Entry(wins,width=30,relief="solid",bd=1,textvariable=a).grid(row=c,column=2,pady=10,padx=10)
                qte=Label(wins,font = ("Times new roman",13,"bold"),text =("B).")).grid(row=c,column=3,pady=10)
                b="opnb"+str(i)
                b=StringVar()

                enh=Entry(wins,width=30,relief="solid",bd=1,textvariable=b).grid(row=c,column=4,pady=10,padx=10)
                c=c+1
                qtd=Label(wins,font = ("Times new roman",13,"bold"),text =("C).")).grid(row=c,column=1)
                w="opnc"+str(i)
                w=StringVar()

                end=Entry(wins,width=30,relief="solid",bd=1,textvariable=w).grid(row=c,column=2)

                qts=Label(wins,text =("D)."),font = ("Times new roman",13,"bold")).grid(row=c,column=3)
                d="opnd"+str(i)
                d=StringVar()

                ens=Entry(wins,width=30,relief="solid",bd=1,textvariable=d).grid(row=c,column=4)
               
                c=c+1
                btr=Button(wins,text="submit",width=20,relief="ridge",bd=3,command=saveme,bg="red",fg="white").grid(row=150,column=2,padx=10,pady=20)
                nw=str(i)+"oprnd"
                nw=[]

                nw.append(k)
                nw.append(a)
                nw.append(b)
                nw.append(w)
                nw.append(d)
                lis.append(nw)   

            print(lis)

        winr=Frame(chos)
        myframe=Frame(chos,relief=GROOVE,width=500,height=100,bd=1)
        # myframe.place(x=10,y=10)
        for box in (winr,myframe):
            box.grid(row=0,column=0,sticky="nsew")


        lad=Label(winr,text="MCQ'S QUESTION MAKER",font=("Timens new roman",15,"bold"),fg="red").grid(row=0,column=3,columnspan=100)
        nwe=IntVar()
        lad=Label(winr,text="NO OF BITS TO CREATE : ",font=("Timens new roman",12,"bold")).grid(row=2,column=2,pady=20,padx=10)
        rntryo=Entry(winr,width=80,textvariable=nwe).grid(row=2,column=3,pady=10)
        nwr=IntVar()
        lade=Label(winr,text="MARKS FOR EACH BIT : ",font=("Timens new roman",12,"bold")).grid(row=3,column=2,pady=10,padx=10)
        rntrry=Entry(winr,width=80,textvariable=nwr).grid(row=3,column=3,pady=10)
        sub=Button(winr,text="Submit",width=20,relief="groove",font=("Timens new roman",12,"bold"),fg="white",bg="red",command=lambda :[next(myframe),done()]).grid(row=4,column=3,pady=20,padx=50)
        
        canvas=Canvas(myframe)
        frame=Frame(canvas)
        myscrollbar=Scrollbar(myframe,orient="vertical",command=canvas.yview)
        canvas.configure(yscrollcommand=myscrollbar.set)

        myscrollbar.pack(side="right",fill="y")
        canvas.pack(side="left")
        canvas.create_window((0,0),window=frame,anchor='nw')
        frame.bind("<Configure>",myfunction)
        next(winr)


        chos.mainloop()
        
        # if(mcq.get()=="Blanks"):
        #     import bitsofflinemaker
        #     bitsofflinemaker.bitmemore()

        # if(mcq.get()=="MCQ's"):
        #     import choosebit
        #     choosebit.notify()

    def bitmemore():  
          

        def data():
            
            y=koi.get()
            p=int(y)
            q=kois.get()
            z=int(q)
            ans=p*z
            m=str(ans)
            s=0

            maxo="MAX MARKS : "+m+" M "
            lis=[]
            # doc=docx.Document()
            # sections = doc.sections
            # for section in sections:
            #     section.top_margin = Cm(1)
            #     section.bottom_margin = Cm(0.5)
            #     section.left_margin = Cm(1.5)
            #     section.right_margin = Cm(1.5)


            print(ans)
            







            def getbit():
                header_text=doc.add_heading("I. Fill in the BLANKS with appropriate Answer : ")
                dft=doc.add_heading(maxo,2)



                dft.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
                for i in lis:
                    print(i.get())
                    para=doc.add_paragraph("                          ")
                    para=doc.add_paragraph(i.get())
                    para.style="List Number"
                saveas.config(state=NORMAL)
                root.destroy()   
                 
                # doc.save("sivas.docx")

            for i in range(1,p+1):
                k="text"+str(i)

                print(k)
                l=Label(frame,text="BIT {} :".format(i)).grid(row=s,column=0,pady=5)
                s=s+1
                k=StringVar()
                e1=Entry(frame,width=80,textvariable=k).grid(row=s,column=0,columnspan=500)
                lis.append(k)
                s=s+1
            b=Button(frame,text="Done",width=10,command=getbit).grid(row=150,column=2,pady=10)


        def myfunction(event):
            canvas.configure(scrollregion=canvas.bbox("all"),width=560,height=380)

        root=Toplevel()
        sizex = 630
        sizey = 600
        posx  = 100
        posy  = 100
        root.wm_geometry("%dx%d+%d+%d" % (sizex, sizey, posx, posy))
        root.geometry("630x600")
        l=Label(root,text="Offline Bit Paper MAKER",font="camble 15 bold").grid(row=0,column=0,columnspan=500,pady=10)
        l=Label(root,text="Enter count of total no of BITS :").grid(row=1,column=0,pady=7)
        koi=IntVar()
        e=Entry(root,width=30,textvariable=koi).grid(row=1,column=1,padx=30)
        l=Label(root,text="Marks of each BITS :").grid(row=2,column=0,pady=20)
        kois=IntVar()
        e1=Entry(root,width=30,textvariable=kois).grid(row=2,column=1,padx=30)
        b1=Button(root,text="Submit",width=10,command=data,bg="red",fg="white",font="camble 13 bold").grid(row=3,column=1,padx=50)
        l=Label(root,text="-------------------------------------------------------------------------------------------------------------------------------").grid(row=4,column=0,columnspan=500)

        myframe=Frame(root,relief=SOLID,width=50,height=100,bd=1)
        myframe.place(x=10,y=200)

        canvas=Canvas(myframe)
        frame=Frame(canvas)
        myscrollbar=Scrollbar(myframe,orient="vertical",command=canvas.yview)
        canvas.configure(yscrollcommand=myscrollbar.set)

        myscrollbar.pack(side="right",fill="y")
        canvas.pack(side="left")
        canvas.create_window((0,0),window=frame,anchor='nw')
        frame.bind("<Configure>",myfunction)
        data()
        root.mainloop()       



    l15=Label(plan,text="Select the below Option").place(x=400,y=390)
    # mcq = StringVar()
    # com = ttk.Combobox(plan, width = 30,  textvariable = mcq) 
    # com['values'] = ("MCQ's",'Blanks','Both','none') 
    # com.place(x=390,y=420)
    # com.current(0)
    btn2=Button(plan,text="BLANKS",fg="red",bg="white",border=2,font='camble 10 bold',command=bitmemore).place(x=410,y=420)
    btn2=Button(plan,text="MCQ'S",fg="red",bg="white",border=2,font='camble 10 bold',command=check).place(x=490,y=420)

    b1["state"] = DISABLED
    

background_image=PhotoImage("unnamed.jpg")
background_label = Label(plan, image=background_image)
background_label.place(x=0, y=0, relwidth=1, relheight=1)



l=Label(plan,text="QUESTION PAPER",font="camble 15 bold",fg='red').pack()
l=Label(plan,text="College/School/InstitutionName :",font="arial 10 bold").place(x=10,y=50)
cn=StringVar()
e1=Entry(plan,width=100,textvariable=cn,relief="ridge",border=3).place(x=230,y=50)
l=Label(plan,text="Subject Code :",font="arial 10 bold").place(x=10,y=100)
sc=StringVar()
e2=Entry(plan,width=40,textvariable=sc,relief="ridge",border=3).place(x=130,y=100)
l=Label(plan,text="Subject Name :",font="arial 10 bold").place(x=490,y=100)
sn=StringVar()
e3=Entry(plan,width=40,textvariable=sn,relief="ridge",border=3).place(x=600,y=100)
l=Label(plan,text="MAX.MARKS :",font="arial 10 bold").place(x=10,y=150)
mm=StringVar()
e4=Entry(plan,width=40,textvariable=mm,relief="ridge",border=3).place(x=130,y=150)
l=Label(plan,text="Duration :",font="arial 10 bold").place(x=490,y=150)
dr=StringVar()
e5=Entry(plan,width=40,textvariable=dr,relief="ridge",border=3).place(x=600,y=150)
l=Label(plan,text="Class / Branch :",font="arial 10 bold").place(x=10,y=200)
cl=StringVar()
e6=Entry(plan,width=40,textvariable=cl,relief="ridge",border=3).place(x=130,y=200)
l=Label(plan,text="DATE/YEAR :",font="arial 10 bold").place(x=490,y=200)
dy=StringVar()
e7=Entry(plan,width=40,textvariable=dy,relief="ridge",border=3).place(x=600,y=200)
l=Label(plan,text="Exam-Type:",font="arial 10 bold").place(x=10,y=250)
n = StringVar() 
com = ttk.Combobox(plan, width = 30,  textvariable = n) 
com['values'] = ('SemisterEndexam','Test','Mid-Term','Anual-Exam','Slip-test','Unit-test','Half-Yearly Exams','Quarterly Exams') 
com.place(x=130,y=250)
com.current(1)
bb=Button(plan,text="Verify and Submit",command=button1,fg="red",bg="white").place(x=730,y=250)
l=Label(plan,text="--------------------------------------------------------------------------------------------------------------------------------------------------------------------------------------------------------------------------------------").place(x=0,y=310)  
l=Label(plan,text="Type of EXAM :").place(x=20,y=310)

b1=Button(plan,text="Discriptive Type",fg="red",bg="white",border=2,width=15,font='camble 10 bold',command=discriptive)
b1.place(x=150,y=350)
b1.config(state=DISABLED)

b2=Button(plan,text="Multiple Choice/Bits Type",fg="white",bg="red",border=2,width=20,font='camble 10 bold',command=mcqs)
b2.place(x=400,y=350)
b2.config(state=DISABLED)

b3=Button(plan,text="Refresh",fg="black",bg="white",width=10,command=refresh)
b3.place(x=700,y=350)

saveas=Button(plan,text="Save & Generate Paper",fg="red",bg="white",width=20,command=saveas)
saveas.config(state=DISABLED)
saveas.place(x=700,y=400)







plan.mainloop()
